from pdf2image import convert_from_path
import cv2
import numpy as np
import pytesseract
import pandas as pd
from multiprocessing import Pool, cpu_count
import logging
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import fonts
from reportlab.pdfgen import canvas
import os
from docx import Document
import comtypes.client

# Especifica la ruta completa al ejecutable de Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'

# Configuración básica de logging
logging.basicConfig(filename='pdf_processing.log', level=logging.ERROR)
    
def convert_pdf_to_images(pdf_file):
    """
    Converts a PDF file into a list of images (one per page).
    """
    return convert_from_path(pdf_file, poppler_path=r'C:/Program Files/poppler-24.07.0/Library/bin')

def deskew(image):
    """
    Corrects the skew of the given image.
    """
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
    gray = cv2.bitwise_not(gray)
    coords = np.column_stack(np.where(gray > 0))
    angle = cv2.minAreaRect(coords)[-1]
    
    if abs(angle) < 0.5:  # If the angle is less than 0.5 degrees, skip deskew
        return image

    if angle < -45:
        angle = -(90 + angle)
    else:
        angle = -angle

    (h, w) = image.shape[:2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    rotated = cv2.warpAffine(image, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

    return rotated

def extract_text_from_image(image):
    """
    Extracts text from a given image using Tesseract OCR.
    """
    return pytesseract.image_to_string(image)

def get_text_data_from_image(image):
    """
    Extracts detailed text data including confidence, and removes headers and footers.
    """
    d = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
    d_df = pd.DataFrame.from_dict(d)
    
    # Calculate block numbers
    block_num = int(d_df.loc[d_df['level'] == 2, 'block_num'].max())
    
    # Identify header and footer index
    header_index = d_df[d_df['block_num'] == 1].index.values
    footer_index = d_df[d_df['block_num'] == block_num].index.values
    
    # Combine text, excluding header and footer
    text = ' '.join(d_df.loc[(d_df['level'] == 5) & (~d_df.index.isin(header_index) & ~d_df.index.isin(footer_index)), 'text'].values)
    
    return text

def process_page(page):
    """
    Processes a single page image: deskews it and extracts text data.
    """
    try:
        page_arr = np.array(page)
        page_arr_gray = cv2.cvtColor(page_arr, cv2.COLOR_BGR2GRAY)
        page_deskewed = deskew(page_arr_gray)
        text_data = get_text_data_from_image(page_deskewed)
        return text_data
    except Exception as e:
        logging.error(f"Error processing page: {str(e)}")
        return f"Error processing page: {str(e)}"
    

def extract_text_from_pdf(pdf_file):
    """
    Extracts text from each page of the given PDF file.
    """
    pages = convert_pdf_to_images(pdf_file)
    extracted_text = []

    with Pool(cpu_count()) as pool:
        extracted_text = pool.map(process_page, pages)

    return extracted_text

def save_to_word(text_from_pdf, output_filename):
    """
    Save the extracted text to a Word document.
    """
    doc = Document()
    for page_num, text in enumerate(text_from_pdf, start=1):
        doc.add_heading(f'Page {page_num}', level=1)
        doc.add_paragraph(text)
        doc.add_page_break()
    doc.save(output_filename)

def convert_word_to_pdf(word_filename, pdf_filename):
    """
    Convert a Word document to PDF using comtypes.
    """
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(word_filename)
    doc.SaveAs(pdf_filename, FileFormat=17)  # FileFormat=17 es para PDF
    doc.Close()
    word.Quit()

if __name__ == "__main__":
    # Replace 'input_file.pdf' with the path to your PDF file
    pdf_name = 'CORPORACIÓN HOGAR BELEN'
    pdf_file = pdf_name + '.pdf'
    text_from_pdf = extract_text_from_pdf(pdf_file)

    # Save the extracted text to a Word document
    pdf_filename = pdf_name + ' word_salida.docx'
    save_to_word(text_from_pdf, pdf_filename)

    # Convert to a PDF
    pdf_filename = 'output_word.pdf'
    convert_word_to_pdf(os.path.abspath(pdf_filename), os.path.abspath(pdf_filename))

    print("Text extraction complete. Word and PDF files have been saved.")

    # You can now save or print the extracted text
    for page_num, text in enumerate(text_from_pdf, start=1):
        print(f"Page {page_num}:")
        print(text)
        print("-" * 50)
